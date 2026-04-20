"""Document processing utilities for AgentMind."""

from enum import Enum
from pathlib import Path
from typing import Optional, Union, List, Dict

try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentFormat(str, Enum):
    """Supported document formats."""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"


class DocumentProcessor:
    """Process documents for multi-modal agent interactions."""

    def __init__(self):
        """Initialize document processor."""

    def extract_text_from_pdf(
        self, path: Union[str, Path], pages: Optional[List[int]] = None
    ) -> str:
        """Extract text from PDF file.

        Args:
            path: Path to PDF file
            pages: List of page numbers to extract (0-indexed), None for all

        Returns:
            Extracted text
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError(
                "PyPDF2 is required for PDF processing. " "Install with: pip install PyPDF2"
            )

        text_parts = []
        with open(path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            total_pages = len(reader.pages)

            if pages is None:
                pages = range(total_pages)

            for page_num in pages:
                if 0 <= page_num < total_pages:
                    page = reader.pages[page_num]
                    text_parts.append(page.extract_text())

        return "\n\n".join(text_parts)

    def extract_text_from_docx(self, path: Union[str, Path]) -> str:
        """Extract text from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument(path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs]
        return "\n\n".join(text_parts)

    def extract_text_from_txt(self, path: Union[str, Path]) -> str:
        """Extract text from TXT file.

        Args:
            path: Path to TXT file

        Returns:
            File contents
        """
        with open(path, "r", encoding="utf-8") as file:
            return file.read()

    def extract_text(self, path: Union[str, Path], format: Optional[DocumentFormat] = None) -> str:
        """Extract text from document (auto-detect format).

        Args:
            path: Path to document
            format: Document format, auto-detected if None

        Returns:
            Extracted text
        """
        path = Path(path)

        if format is None:
            ext = path.suffix.lower().lstrip(".")
            if ext in DocumentFormat.__members__.values():
                format = DocumentFormat(ext)
            else:
                raise ValueError(f"Unsupported file extension: {ext}")

        if format == DocumentFormat.PDF:
            return self.extract_text_from_pdf(path)
        elif format == DocumentFormat.DOCX:
            return self.extract_text_from_docx(path)
        elif format in (DocumentFormat.TXT, DocumentFormat.MD):
            return self.extract_text_from_txt(path)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def get_pdf_metadata(self, path: Union[str, Path]) -> Dict:
        """Get metadata from PDF file.

        Args:
            path: Path to PDF file

        Returns:
            Dictionary with PDF metadata
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError(
                "PyPDF2 is required for PDF processing. " "Install with: pip install PyPDF2"
            )

        with open(path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            metadata = reader.metadata or {}

            return {
                "title": metadata.get("/Title", ""),
                "author": metadata.get("/Author", ""),
                "subject": metadata.get("/Subject", ""),
                "creator": metadata.get("/Creator", ""),
                "producer": metadata.get("/Producer", ""),
                "num_pages": len(reader.pages),
            }

    def get_docx_metadata(self, path: Union[str, Path]) -> Dict:
        """Get metadata from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            Dictionary with DOCX metadata
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument(path)
        core_props = doc.core_properties

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": core_props.created,
            "modified": core_props.modified,
            "num_paragraphs": len(doc.paragraphs),
        }

    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks.

        Args:
            text: Text to chunk
            chunk_size: Maximum chunk size in characters
            overlap: Overlap between chunks in characters

        Returns:
            List of text chunks
        """
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap

        return chunks

    def prepare_for_llm(
        self,
        path: Union[str, Path],
        format: Optional[DocumentFormat] = None,
        chunk_size: Optional[int] = None,
        pages: Optional[List[int]] = None,
    ) -> dict:
        """Prepare document for LLM input.

        Args:
            path: Path to document
            format: Document format, auto-detected if None
            chunk_size: If provided, split text into chunks
            pages: For PDFs, specific pages to extract

        Returns:
            Dictionary with document data for LLM
        """
        # Extract text
        if format == DocumentFormat.PDF and pages is not None:
            text = self.extract_text_from_pdf(path, pages=pages)
        else:
            text = self.extract_text(path, format=format)

        # Get metadata
        path = Path(path)
        if format is None:
            ext = path.suffix.lower().lstrip(".")
            format = DocumentFormat(ext) if ext in DocumentFormat.__members__.values() else None

        metadata = {}
        if format == DocumentFormat.PDF:
            try:
                metadata = self.get_pdf_metadata(path)
            except Exception:
                pass
        elif format == DocumentFormat.DOCX:
            try:
                metadata = self.get_docx_metadata(path)
            except Exception:
                pass

        result = {
            "type": "document",
            "format": format.value if format else "unknown",
            "filename": path.name,
            "text": text,
            "metadata": metadata,
        }

        # Chunk if requested
        if chunk_size:
            result["chunks"] = self.chunk_text(text, chunk_size=chunk_size)

        return result

    def save_text_to_file(
        self, text: str, path: Union[str, Path], format: DocumentFormat = DocumentFormat.TXT
    ) -> None:
        """Save text to file.

        Args:
            text: Text to save
            path: Output file path
            format: Output format
        """
        if format in (DocumentFormat.TXT, DocumentFormat.MD):
            with open(path, "w", encoding="utf-8") as file:
                file.write(text)
        else:
            raise ValueError(f"Saving to {format} format is not supported")
